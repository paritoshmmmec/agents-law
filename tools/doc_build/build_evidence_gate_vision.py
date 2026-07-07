from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "docs" / "Evidence_Gate_Ideal_Vision.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(11, 37, 69)
MUTED = RGBColor(89, 89, 89)
CALLOUT_FILL = "F4F6F9"
CALLOUT_BORDER = "D9E2EC"


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_border_bottom(paragraph, color="2E74B5", size="8", space="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def shade_paragraph(paragraph, fill=CALLOUT_FILL, border=CALLOUT_BORDER):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    for side in ("top", "left", "bottom", "right"):
        edge = OxmlElement(f"w:{side}")
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), "4")
        edge.set(qn("w:space"), "4")
        edge.set(qn("w:color"), border)
        borders.append(edge)
    p_pr.append(borders)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.10

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167


def add_header_footer(doc):
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.text = ""
    p.paragraph_format.space_after = Pt(0)
    left = p.add_run("Evidence Gate")
    set_run_font(left, size=9, color=MUTED, bold=True)
    p.add_run("\t")
    right = p.add_run("Ideal Vision")
    set_run_font(right, size=9, color=MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.text = ""
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(0)
    run = fp.add_run("Vision brief")
    set_run_font(run, size=9, color=MUTED)


def add_masthead(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("EVIDENCE GATE")
    set_run_font(r, size=23, color=INK, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(16)
    r = p.add_run("Ideal Vision for the Deterministic Evidence Firewall for AI Agents")
    set_run_font(r, size=14, color=MUTED)

    metadata = [
        ("Type", "Product vision brief"),
        ("Date", "July 7, 2026"),
        ("Core idea", "AI agents should prove their work before they act."),
        ("Audience", "Agent builders, security teams, platform teams, and compliance owners"),
    ]
    for label, value in metadata:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        lr = p.add_run(f"{label}: ")
        set_run_font(lr, size=11, color=INK, bold=True)
        vr = p.add_run(value)
        set_run_font(vr, size=11, color=INK)

    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(14)
    rule.paragraph_format.space_after = Pt(12)
    set_paragraph_border_bottom(rule)


def add_body_paragraph(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, size=11, color=RGBColor(0, 0, 0))


def add_callout(doc, label, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(10)
    shade_paragraph(p)
    lr = p.add_run(f"{label}: ")
    set_run_font(lr, size=11, color=INK, bold=True)
    tr = p.add_run(text)
    set_run_font(tr, size=11, color=INK)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run(item)
        set_run_font(run, size=11, color=RGBColor(0, 0, 0))


def build():
    doc = Document()
    configure_document(doc)
    add_header_footer(doc)
    add_masthead(doc)

    doc.add_heading("Ideal Vision", level=1)
    add_body_paragraph(
        doc,
        "Evidence Gate becomes the trust layer for agentic software: a small, "
        "deterministic enforcement system that sits between an AI agent and any "
        "consequential action, requiring the agent to prove that its action is "
        "justified by fresh, authorized, non-conflicting evidence before anything executes.",
    )
    add_body_paragraph(
        doc,
        "A team points Evidence Gate at yesterday's agent traces and immediately "
        "sees which actions would have been allowed, restricted, sent to review, "
        "or blocked. No production risk, no rewrite. From there, they wrap one "
        "sensitive tool by name pattern, connect the same policy pack, and move "
        "from simulation to live fail-closed enforcement.",
    )
    add_callout(
        doc,
        "Product promise",
        "Evidence-based authorization for agents: RBAC decides whether an agent "
        "may use a tool; Evidence Gate decides whether this specific tool call is "
        "justified right now.",
    )

    doc.add_heading("North Star", level=1)
    add_body_paragraph(
        doc,
        "Make it easy for any team to answer: Can we prove why this agent took "
        "this action, using facts it was actually allowed to rely on?"
    )
    add_body_paragraph(
        doc,
        "If the answer is no, the system blocks, restricts, or routes to review "
        "before harm occurs."
    )

    doc.add_heading("What It Should Become", level=1)
    add_bullets(
        doc,
        [
            "Deterministic at runtime: no LLM decides enforcement, and every decision is reproducible.",
            "Easy to adopt: start with trace replay, then wrap existing tools with minimal code changes.",
            "Hard to bypass: sensitive tools require clearance, and downstream services can verify signed tokens.",
            "Audit-native: every decision is hash-chained, replayable, and explainable without exposing raw prompts or sensitive payloads.",
            "Framework-agnostic: works with LangChain, CrewAI, LlamaIndex, custom agents, remote services, and direct Python functions.",
            "Policy-readable: humans can inspect and approve rule packs; LLMs may help draft policies, but never silently enforce them.",
        ],
    )

    doc.add_heading("Ideal Positioning", level=1)
    add_body_paragraph(
        doc,
        "Evidence Gate is not an agent framework, observability tool, or RBAC replacement. "
        "It is the deterministic evidence firewall for AI agents."
    )
    add_callout(
        doc,
        "Positioning line",
        "The agent is authorized, but is this action actually justified by the evidence?",
    )

    doc.add_heading("Ideal Product Flow", level=1)
    steps = [
        "Replay: import LangSmith, OpenAI, Langfuse, or custom traces.",
        "Diagnose: see which actions would be allowed, reviewed, restricted, or blocked, with reasons.",
        "Author: write or generate readable YAML policies from SOPs, then approve them.",
        "Enforce: wrap sensitive tools by name pattern or framework adapter.",
        "Verify: downstream tools require signed clearance tokens.",
        "Audit: export replayable, tamper-evident decision logs for security, compliance, and debugging.",
    ]
    for step in steps:
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(step)
        set_run_font(run, size=11, color=RGBColor(0, 0, 0))

    doc.add_heading("Ideal One-Liners", level=1)
    add_callout(doc, "Plain English", "Evidence Gate makes AI agents prove their work before they act.")
    add_callout(
        doc,
        "Enterprise",
        "Evidence Gate is a deterministic runtime control plane for agentic systems, "
        "enforcing evidence-based policies before sensitive tool calls execute and "
        "producing verifiable audit trails for every decision.",
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
